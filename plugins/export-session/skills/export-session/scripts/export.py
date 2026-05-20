#!/usr/bin/env python3
"""
export.py — Claude Code session export script
Converts ~/.claude/projects JSONL transcripts to md, json, txt, docx, or pdf.
"""

import argparse
import json
import os
import sys
from datetime import datetime, timezone
from pathlib import Path


def parse_args():
    p = argparse.ArgumentParser()
    p.add_argument("--input", required=True, help="Path to session .jsonl file")
    p.add_argument("--output", required=True, help="Output file path")
    p.add_argument("--format", required=True, choices=["md", "json", "txt", "docx", "pdf"])
    p.add_argument("--last", type=int, default=0, help="Export last N assistant turns (0 = all)")
    return p.parse_args()


def load_turns(jsonl_path: str, last: int) -> list[dict]:
    """Load assistant turns from JSONL, optionally limited to last N."""
    turns = []
    with open(jsonl_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                entry = json.loads(line)
            except json.JSONDecodeError:
                continue

            # Handle both flat and nested message formats
            role = entry.get("role") or entry.get("type")
            msg = entry.get("message", entry)
            if msg.get("role") == "assistant" or role == "assistant":
                content = extract_content(msg)
                if content:
                    turns.append({
                        "role": "assistant",
                        "content": content,
                        "timestamp": entry.get("timestamp", ""),
                    })

    if last and last > 0:
        turns = turns[-last:]
    return turns


def extract_content(msg: dict) -> str:
    """Extract plain text content from a message object."""
    content = msg.get("content", "")
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts = []
        for block in content:
            if isinstance(block, dict) and block.get("type") == "text":
                parts.append(block.get("text", ""))
        return "\n".join(parts)
    return str(content)


def to_md(turns: list[dict], session_path: str) -> str:
    date = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    lines = [f"# Session Export — {date}", f"", f"*Source: `{session_path}`*", ""]
    for i, t in enumerate(turns, 1):
        lines.append(f"## Response {i}")
        if t["timestamp"]:
            lines.append(f"*{t['timestamp']}*")
            lines.append("")
        lines.append(t["content"])
        lines.append("")
        lines.append("---")
        lines.append("")
    return "\n".join(lines)


def to_txt(turns: list[dict]) -> str:
    parts = []
    for i, t in enumerate(turns, 1):
        parts.append(f"=== Response {i} ===")
        if t["timestamp"]:
            parts.append(t["timestamp"])
        parts.append(t["content"])
        parts.append("")
    return "\n".join(parts)


def to_json(turns: list[dict], session_path: str) -> str:
    payload = {
        "exported_at": datetime.now(timezone.utc).isoformat(),
        "source": session_path,
        "turn_count": len(turns),
        "turns": turns,
    }
    return json.dumps(payload, indent=2, ensure_ascii=False)


def to_docx(turns: list[dict], output_path: str, session_path: str):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx --break-system-packages", file=sys.stderr)
        sys.exit(1)

    doc = Document()

    # Title
    title = doc.add_heading("Session Export", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata
    date = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    meta = doc.add_paragraph(f"Exported: {date}\nSource: {session_path}")
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph("")

    for i, t in enumerate(turns, 1):
        heading = doc.add_heading(f"Response {i}", level=2)
        if t["timestamp"]:
            ts = doc.add_paragraph(t["timestamp"])
            ts.runs[0].font.size = Pt(9)
            ts.runs[0].italic = True

        # Split on double newlines to preserve paragraph breaks
        for para in t["content"].split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        doc.add_paragraph("─" * 40)

    doc.save(output_path)


def to_pdf(turns: list[dict], output_path: str, session_path: str):
    try:
        from fpdf import FPDF
    except ImportError:
        print("ERROR: fpdf2 not installed. Run: pip install fpdf2 --break-system-packages", file=sys.stderr)
        sys.exit(1)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Session Export", ln=True)

    # Metadata
    date = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 6, f"Exported: {date}", ln=True)
    pdf.cell(0, 6, f"Source: {os.path.basename(session_path)}", ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    for i, t in enumerate(turns, 1):
        # Response heading
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 8, f"Response {i}", ln=True)

        if t["timestamp"]:
            pdf.set_font("Helvetica", "I", 9)
            pdf.set_text_color(100, 100, 100)
            pdf.cell(0, 5, t["timestamp"], ln=True)
            pdf.set_text_color(0, 0, 0)

        # Content — fpdf2 multi_cell handles line wrapping
        pdf.set_font("Helvetica", "", 10)
        # Sanitize: fpdf latin-1 safe (replace non-latin chars)
        safe = t["content"].encode("latin-1", errors="replace").decode("latin-1")
        pdf.multi_cell(0, 5, safe)
        pdf.ln(2)

        # Divider
        pdf.set_draw_color(200, 200, 200)
        pdf.line(pdf.get_x(), pdf.get_y(), 200, pdf.get_y())
        pdf.ln(4)

    pdf.output(output_path)


def main():
    args = parse_args()
    output_path = str(Path(args.output).expanduser().resolve())
    session_path = str(Path(args.input).expanduser().resolve())

    if not os.path.exists(session_path):
        print(f"ERROR: JSONL not found: {session_path}", file=sys.stderr)
        sys.exit(1)

    turns = load_turns(session_path, args.last)
    if not turns:
        print("WARNING: No assistant turns found in session.", file=sys.stderr)
        sys.exit(1)

    fmt = args.format
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

    if fmt == "md":
        Path(output_path).write_text(to_md(turns, session_path), encoding="utf-8")
    elif fmt == "txt":
        Path(output_path).write_text(to_txt(turns), encoding="utf-8")
    elif fmt == "json":
        Path(output_path).write_text(to_json(turns, session_path), encoding="utf-8")
    elif fmt == "docx":
        to_docx(turns, output_path, session_path)
    elif fmt == "pdf":
        to_pdf(turns, output_path, session_path)

    print(f"✓ Exported {len(turns)} turn(s) to: {output_path}")


if __name__ == "__main__":
    main()